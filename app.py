#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PROT-Relatório-Automático: Sistema de geração automática de relatórios de entrega de proteínas
Autor: Sistema PROT
Data: 2026-07-09
"""

from flask import Flask, render_template, request, jsonify, send_file
from werkzeug.utils import secure_filename
import sqlite3
import os
import json
from datetime import datetime
from PIL import Image
import pytesseract
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import traceback

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['DATABASE'] = 'prot_database.db'

# Criar pasta de uploads
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# ==================== BANCO DE DADOS ====================

def init_db():
      """Inicializa o banco de dados SQLite"""
      conn = sqlite3.connect(app.config['DATABASE'])
      c = conn.cursor()

    c.execute('''CREATE TABLE IF NOT EXISTS entregas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
                    num_prot TEXT UNIQUE NOT NULL,
                            data_entrega TEXT NOT NULL,
                                    nf_numero TEXT NOT NULL,
                                            nf_serie TEXT NOT NULL,
                                                    nf_emissor TEXT NOT NULL,
                                                            nf_data TEXT NOT NULL,
                                                                    nf_valor REAL NOT NULL,
                                                                            classificacao TEXT DEFAULT 'CONFORME',
                                                                                    observacoes TEXT,
                                                                                            criado_em TEXT DEFAULT CURRENT_TIMESTAMP
                                                                                                )''')

    c.execute('''CREATE TABLE IF NOT EXISTS proteinas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
                    entrega_id INTEGER NOT NULL,
                            tipo TEXT NOT NULL,
                                    peso_kg REAL NOT NULL,
                                            lote TEXT NOT NULL,
                                                    data_fabricacao TEXT NOT NULL,
                                                            data_validade TEXT NOT NULL,
                                                                    quantidade_unidades INTEGER,
                                                                            FOREIGN KEY (entrega_id) REFERENCES entregas(id)
                                                                                )''')

    c.execute('''CREATE TABLE IF NOT EXISTS fotos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
                    entrega_id INTEGER NOT NULL,
                            tipo_foto TEXT NOT NULL,
                                    caminho_arquivo TEXT NOT NULL,
                                            criada_em TEXT DEFAULT CURRENT_TIMESTAMP,
                                                    FOREIGN KEY (entrega_id) REFERENCES entregas(id)
                                                        )''')

    conn.commit()
    conn.close()

init_db()

# ==================== FUNÇÕES AUXILIARES ====================

def extrair_dados_imagem(caminho_imagem):
      """Extrai texto de uma imagem usando OCR (Tesseract)"""
      try:
                img = Image.open(caminho_imagem)
                texto = pytesseract.image_to_string(img, lang='por')
                return texto
except Exception as e:
        return f"Erro ao processar imagem: {str(e)}"

def salvar_entrega(num_prot, data_entrega, nf_info, proteinas_list):
      """Salva uma entrega no banco de dados"""
      conn = sqlite3.connect(app.config['DATABASE'])
      c = conn.cursor()

    try:
              c.execute('''INSERT INTO entregas 
                                   (num_prot, data_entrega, nf_numero, nf_serie, nf_emissor, nf_data, nf_valor)
                                                        VALUES (?, ?, ?, ?, ?, ?, ?)''',
                                          (num_prot, data_entrega, nf_info['numero'], nf_info['serie'], 
                                           nf_info['emissor'], nf_info['data'], nf_info['valor']))

        entrega_id = c.lastrowid

        # Salvar proteínas
        for prot in proteinas_list:
                      c.execute('''INSERT INTO proteinas 
                                               (entrega_id, tipo, peso_kg, lote, data_fabricacao, data_validade)
                                                                        VALUES (?, ?, ?, ?, ?, ?)''',
                                                      (entrega_id, prot['tipo'], prot['peso'], prot['lote'], 
                                                       prot['fabricacao'], prot['validade']))

        conn.commit()
        return entrega_id
except sqlite3.IntegrityError:
        return None
finally:
        conn.close()

def gerar_pdf_relatorio(entrega_id):
      """Gera relatório em PDF"""
      conn = sqlite3.connect(app.config['DATABASE'])
      c = conn.cursor()

    # Buscar dados da entrega
      c.execute('SELECT * FROM entregas WHERE id = ?', (entrega_id,))
    entrega = c.fetchone()

    c.execute('SELECT * FROM proteinas WHERE entrega_id = ?', (entrega_id,))
    proteinas = c.fetchall()

    conn.close()

    if not entrega:
              return None

    # Criar PDF
    filename = f"PROT-{entrega[1]}.pdf"
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)

    doc = SimpleDocTemplate(filepath, pagesize=A4, topMargin=0.5*inch, bottomMargin=0.5*inch)
    story = []
    styles = getSampleStyleSheet()

    # Título
    title_style = ParagraphStyle(
              'CustomTitle',
              parent=styles['Heading1'],
              fontSize=16,
              textColor=colors.HexColor('#1B5E20'),
              spaceAfter=12,
              alignment=1
    )
    story.append(Paragraph(f"RELATÓRIO DE ENTREGA DE PROTEÍNAS - {entrega[1]}", title_style))
    story.append(Spacer(1, 0.2*inch))

    # Informações básicas
    info_style = styles['Normal']
    story.append(Paragraph(f"<b>Data da Entrega:</b> {entrega[2]}", info_style))
    story.append(Paragraph(f"<b>NF-e:</b> {entrega[3]}.{entrega[4]} | <b>Emissor:</b> {entrega[5]}", info_style))
    story.append(Paragraph(f"<b>Valor Total:</b> R$ {entrega[7]:.2f}", info_style))
    story.append(Spacer(1, 0.2*inch))

    # Tabela de proteínas
    table_data = [['Tipo de Proteína', 'Peso (kg)', 'Lote', 'Fabricação', 'Validade']]
    for prot in proteinas:
              table_data.append([prot[2], str(prot[3]), prot[4], prot[5], prot[6]])

    table = Table(table_data, colWidths=[1.8*inch, 1.2*inch, 1.2*inch, 1.2*inch, 1.2*inch])
    table.setStyle(TableStyle([
              ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1B5E20')),
              ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
              ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
              ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
              ('FONTSIZE', (0, 0), (-1, 0), 10),
              ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
              ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    story.append(table)
    story.append(Spacer(1, 0.3*inch))

    # Classificação
    classificacao = entrega[8] if len(entrega) > 8 else 'CONFORME'
    status_color = colors.green if classificacao == 'CONFORME' else colors.red
    story.append(Paragraph(f"<b>Classificação:</b> <font color='{status_color.hexValue()}'>{classificacao}</font>", info_style))

    # Build PDF
    doc.build(story)
    return filepath

def gerar_docx_relatorio(entrega_id):
      """Gera relatório em DOCX"""
      conn = sqlite3.connect(app.config['DATABASE'])
      c = conn.cursor()

    c.execute('SELECT * FROM entregas WHERE id = ?', (entrega_id,))
    entrega = c.fetchone()

    c.execute('SELECT * FROM proteinas WHERE entrega_id = ?', (entrega_id,))
    proteinas = c.fetchall()

    conn.close()

    if not entrega:
              return None

    # Criar DOCX
    doc = Document()

    # Título
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"RELATÓRIO DE ENTREGA DE PROTEÍNAS - {entrega[1]}")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(27, 94, 32)

    doc.add_paragraph()

    # Informações
    doc.add_paragraph(f"Data da Entrega: {entrega[2]}")
    doc.add_paragraph(f"NF-e: {entrega[3]}.{entrega[4]} | Emissor: {entrega[5]}")
    doc.add_paragraph(f"Valor Total: R$ {entrega[7]:.2f}")

    doc.add_paragraph()

    # Tabela
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tipo de Proteína'
    hdr_cells[1].text = 'Peso (kg)'
    hdr_cells[2].text = 'Lote'
    hdr_cells[3].text = 'Fabricação'
    hdr_cells[4].text = 'Validade'

    for prot in proteinas:
              row_cells = table.add_row().cells
              row_cells[0].text = prot[2]
              row_cells[1].text = str(prot[3])
              row_cells[2].text = prot[4]
              row_cells[3].text = prot[5]
              row_cells[4].text = prot[6]

    doc.add_paragraph()

    classificacao = entrega[8] if len(entrega) > 8 else 'CONFORME'
    p = doc.add_paragraph(f"Classificação: {classificacao}")
    if classificacao == 'CONFORME':
              p.runs[0].font.color.rgb = RGBColor(0, 128, 0)

    # Salvar
    filename = f"PROT-{entrega[1]}.docx"
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    doc.save(filepath)
    return filepath

# ==================== ROTAS ====================

@app.route('/')
def index():
      """Página inicial"""
      return render_template('index.html')

@app.route('/api/entregas', methods=['GET'])
def listar_entregas():
      """Lista todas as entregas"""
      conn = sqlite3.connect(app.config['DATABASE'])
      c = conn.cursor()
      c.execute('SELECT id, num_prot, data_entrega, nf_numero, classificacao FROM entregas ORDER BY criado_em DESC')
      entregas = c.fetchall()
      conn.close()

    return jsonify([{
              'id': e[0],
              'num_prot': e[1],
              'data_entrega': e[2],
              'nf': e[3],
              'classificacao': e[4]
    } for e in entregas])

@app.route('/api/entrega/<int:entrega_id>', methods=['GET'])
def obter_entrega(entrega_id):
      """Obtém dados de uma entrega específica"""
      conn = sqlite3.connect(app.config['DATABASE'])
      c = conn.cursor()

    c.execute('SELECT * FROM entregas WHERE id = ?', (entrega_id,))
    entrega = c.fetchone()

    if not entrega:
              return jsonify({'erro': 'Entrega não encontrada'}), 404

    c.execute('SELECT * FROM proteinas WHERE entrega_id = ?', (entrega_id,))
    proteinas = c.fetchall()

    conn.close()

    return jsonify({
              'id': entrega[0],
              'num_prot': entrega[1],
              'data_entrega': entrega[2],
              'nf': {
                            'numero': entrega[3],
                            'serie': entrega[4],
                            'emissor': entrega[5],
                            'data': entrega[6],
                            'valor': entrega[7]
              },
              'classificacao': entrega[8],
              'proteinas': [{
                            'tipo': p[2],
                            'peso': p[3],
                            'lote': p[4],
                            'fabricacao': p[5],
                            'validade': p[6]
              } for p in proteinas]
    })

@app.route('/api/criar-relatorio', methods=['POST'])
def criar_relatorio():
      """Cria um novo relatório"""
      try:
                dados = request.json

        # Validar dados
                if not all(k in dados for k in ['num_prot', 'data_entrega', 'nf_info', 'proteinas']):
                              return jsonify({'erro': 'Dados incompletos'}), 400

                # Salvar entrega
                entrega_id = salvar_entrega(
                    dados['num_prot'],
                    dados['data_entrega'],
                    dados['nf_info'],
                    dados['proteinas']
                )

        if not entrega_id:
                      return jsonify({'erro': 'PROT já existe'}), 409

        # Gerar relatórios
        pdf_path = gerar_pdf_relatorio(entrega_id)
        docx_path = gerar_docx_relatorio(entrega_id)

        return jsonify({
                      'sucesso': True,
                      'entrega_id': entrega_id,
                      'num_prot': dados['num_prot'],
                      'pdf': pdf_path.split('/')[-1] if pdf_path else None,
                      'docx': docx_path.split('/')[-1] if docx_path else None
        })
except Exception as e:
        return jsonify({'erro': str(e), 'traceback': traceback.format_exc()}), 500

@app.route('/download/<filename>', methods=['GET'])
def download_arquivo(filename):
      """Download de arquivo"""
    try:
              filepath = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(filename))
              if os.path.exists(filepath):
                            return send_file(filepath, as_attachment=True)
                        return jsonify({'erro': 'Arquivo não encontrado'}), 404
except Exception as e:
        return jsonify({'erro': str(e)}), 500

if __name__ == '__main__':
      app.run(debug=True, host='0.0.0.0', port=5000)
